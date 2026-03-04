"""
Text extraction from various file formats.
Returns raw text for LLM consumption.
"""

import base64
import csv
import io
import json


def extract_text(file_bytes: bytes, content_type: str, filename: str = "") -> str:
    """Extract text content from file bytes based on content type."""
    ct = content_type.lower()

    if ct == "application/pdf":
        return _extract_pdf(file_bytes)
    if ct in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    ):
        return _extract_docx(file_bytes)
    if ct in (
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "application/vnd.ms-excel",
    ):
        return _extract_xlsx(file_bytes)
    if ct == "application/json":
        return _extract_json(file_bytes)
    if ct == "text/csv":
        return _extract_csv(file_bytes)
    if ct.startswith("text/"):
        return _extract_text(file_bytes)
    if ct.startswith("image/"):
        return _extract_image(file_bytes, ct)

    # Fallback: try as text
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext in ("json",):
        return _extract_json(file_bytes)
    if ext in ("csv",):
        return _extract_csv(file_bytes)
    if ext in ("xlsx", "xls"):
        return _extract_xlsx(file_bytes)
    if ext in ("docx", "doc"):
        return _extract_docx(file_bytes)
    if ext in ("pdf",):
        return _extract_pdf(file_bytes)

    return _extract_text(file_bytes)


def _extract_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF using pdfplumber."""
    import pdfplumber

    parts: list[str] = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for i, page in enumerate(pdf.pages):
            text = page.extract_text() or ""
            if text.strip():
                parts.append(f"--- Page {i + 1} ---\n{text}")

            # Extract tables if present
            tables = page.extract_tables()
            for t_idx, table in enumerate(tables):
                if table:
                    parts.append(f"[Tableau page {i + 1}, #{t_idx + 1}]")
                    for row in table:
                        cleaned = [str(cell or "").strip() for cell in row]
                        parts.append(" | ".join(cleaned))

    return "\n\n".join(parts) if parts else "[PDF vide ou illisible]"


def _extract_docx(file_bytes: bytes) -> str:
    """Extract text from Word .docx files."""
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    parts: list[str] = []

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    for table in doc.tables:
        parts.append("[Tableau]")
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            parts.append(" | ".join(cells))

    return "\n".join(parts) if parts else "[Document Word vide]"


def _extract_xlsx(file_bytes: bytes) -> str:
    """Extract text from Excel .xlsx files."""
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
    parts: list[str] = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        parts.append(f"=== Feuille: {sheet_name} ===")
        for row in ws.iter_rows(values_only=True):
            cells = [str(cell if cell is not None else "") for cell in row]
            if any(c.strip() for c in cells):
                parts.append(" | ".join(cells))

    wb.close()
    return "\n".join(parts) if parts else "[Fichier Excel vide]"


def _extract_json(file_bytes: bytes) -> str:
    """Pretty-print JSON content."""
    try:
        data = json.loads(file_bytes.decode("utf-8"))
        return json.dumps(data, indent=2, ensure_ascii=False)
    except (json.JSONDecodeError, UnicodeDecodeError):
        return _extract_text(file_bytes)


def _extract_csv(file_bytes: bytes) -> str:
    """Extract CSV content as pipe-delimited text."""
    text = _decode_bytes(file_bytes)
    parts: list[str] = []
    try:
        dialect = csv.Sniffer().sniff(text[:2048])
        reader = csv.reader(io.StringIO(text), dialect)
    except csv.Error:
        reader = csv.reader(io.StringIO(text))

    for row in reader:
        parts.append(" | ".join(row))

    return "\n".join(parts) if parts else "[CSV vide]"


def _extract_text(file_bytes: bytes) -> str:
    """Extract plain text with encoding detection."""
    return _decode_bytes(file_bytes)


def _extract_image(file_bytes: bytes, content_type: str) -> str:
    """Encode image as base64 for LLM vision analysis."""
    b64 = base64.b64encode(file_bytes).decode("ascii")
    return f"[IMAGE:{content_type};base64,{b64}]"


def _decode_bytes(data: bytes) -> str:
    """Decode bytes trying utf-8 first, then latin-1 fallback."""
    try:
        return data.decode("utf-8")
    except UnicodeDecodeError:
        return data.decode("latin-1", errors="replace")
